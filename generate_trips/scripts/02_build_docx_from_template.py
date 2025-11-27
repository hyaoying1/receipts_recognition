# build_docx_from_templates.py
import os
import json
from copy import deepcopy
from docx import Document
import sys
from pathlib import Path
PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))
from platform_specs import PLATFORM_SPECS


DATA_PATH = PROJECT_ROOT / "data/raw/fake_all.json"
OUT_DOCX_DIR = PROJECT_ROOT /"artifacts/docx"

os.makedirs(OUT_DOCX_DIR, exist_ok=True)

def _cell_full_text(cell):
    return "".join(r.text for p in cell.paragraphs for r in p.runs)

def _find_row_index_by_marker(table, marker):
    """返回 (row_idx, col_idx_of_marker)。找不到抛错。"""
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            if marker in _cell_full_text(cell):
                return ri, ci
    raise ValueError(f"未在表内找到标记 {marker}")

def _append_from_template_row(table, tpl_row):
    new_tr = deepcopy(tpl_row._tr)
    table._tbl.append(new_tr)
    return table.rows[-1]

def _fill_row_by_map(row, amap, trip, idx, receipt):
    for col_idx, fn in amap.items():
        val = fn(trip, idx, receipt)
        set_cell(row.cells[col_idx], "" if val is None else str(val))

def fill_table_row_block(table, trips, spec, receipt):
    """
    针对含 row_block 的平台：
    - 用 detail_marker / summary_marker 定位两条模板行
    - 清掉模板行之下的旧数据
    - 对每个 trip 复制“明细行 + 小结行”
    - 最后删除两条模板行（先删较大索引）
    """
    rb = spec["row_block"]
    d_marker = rb["detail_marker"]
    s_marker = rb["summary_marker"]

    d_idx, d_col = _find_row_index_by_marker(table, d_marker)
    s_idx, s_col = _find_row_index_by_marker(table, s_marker)

    # 统一顺序 & 抓模板行
    top = min(d_idx, s_idx)
    bot = max(d_idx, s_idx)
    d_tpl = table.rows[d_idx]
    s_tpl = table.rows[s_idx]

    # 清掉模板行之下的既有数据（保留表头与两条模板）
    for row in list(table.rows)[bot+1:]:
        table._tbl.remove(row._tr)

    # 逐笔追加“明细 + 小结”
    for idx, trip in enumerate(trips, start=1):
        # 明细行
        new_d = _append_from_template_row(table, d_tpl)
        _fill_row_by_map(new_d, rb["detail_map"], trip, idx, receipt)

        # 小结行：支持 summary_map 或 summary_builder
        new_s = _append_from_template_row(table, s_tpl)
        if "summary_map" in rb:
            _fill_row_by_map(new_s, rb["summary_map"], trip, idx, receipt)
        else:
            # 在小结行里找到 summary_marker 所在的单元格，写入整行文本
            text_cell = new_s.cells[s_col]
            builder = rb["summary_builder"]
            text = builder(trip, idx, receipt)
            set_cell(text_cell, text)

    # 删除模板两行（先删索引大的）
    table._tbl.remove(table.rows[bot]._tr)
    table._tbl.remove(table.rows[top]._tr)

# def fill_header_placeholders(doc, receipt, platform):
#     spec = PLATFORM_SPECS[platform]
#     header_map = spec.get("header_map", {})

#     # 先算好「占位符 -> 替换值」，并统一转成 str，避免 int 报错
#     mapping = {}
#     for ph, fn in header_map.items():
#         try:
#             val = fn(receipt)
#         except Exception:
#             val = ""
#         if val is None:
#             val = ""
#         mapping[ph] = str(val)

#     # 1) 只改包含占位符的 run，不碰别的（比如图片 run）
#     for p in doc.paragraphs:
#         for run in p.runs:
#             original = run.text
#             new_text = original
#             for k, v in mapping.items():
#                 if k in new_text:
#                     new_text = new_text.replace(k, v)
#             if new_text != original:
#                 run.text = new_text  # 只有真的替换了才写回

#     # 2) 表格里的占位符同理
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for p in cell.paragraphs:
#                     for run in p.runs:
#                         original = run.text
#                         new_text = original
#                         for k, v in mapping.items():
#                             if k in new_text:
#                                 new_text = new_text.replace(k, v)
#                         if new_text != original:
#                             run.text = new_text
def fill_header_placeholders(doc, receipt, platform):
    spec = PLATFORM_SPECS[platform]
    header_map = spec.get("header_map", {})

    # 1) 计算占位符 -> 文本（全部转成 str）
    mapping = {}
    for ph, fn in header_map.items():
        try:
            val = fn(receipt)
        except Exception:
            val = ""
        mapping[ph] = "" if val is None else str(val)

    # 工具：段落是否包含图片/绘图（有的话我们跳过，不回写，保护 logo）
    def para_has_drawing(p):
        try:
            return any(r._r.xpath(".//w:drawing") for r in p.runs)
        except Exception:
            return False

    # 工具：对整个段落做一次替换（处理跨 run 的占位符）
    def replace_in_paragraph(p):
        if para_has_drawing(p):
            return  # 跳过含图片的段落

        full = "".join(r.text for r in p.runs)
        if not full:
            return
        changed = False
        for k, v in mapping.items():
            if k in full:
                full = full.replace(k, v)
                changed = True
        if not changed:
            return

        # 只在发生替换时回写；回写成一个 run，沿用第一个 run 的样式
        if p.runs:
            first_style = p.runs[0].style
            for r in p.runs:
                r.text = ""   # 此段落没有图片，安全
            p.runs[0].text = full
            if first_style:
                p.runs[0].style = first_style
        else:
            p.add_run(full)

    # 2) 顶层段落
    for p in doc.paragraphs:
        replace_in_paragraph(p)

    # 3) 表格中的段落
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)


def set_cell(cell, value):
    """
    只改文字，不动单元格样式。
    """
    value = "" if value is None else str(value)

    if not cell.paragraphs:
        p = cell.add_paragraph()
        p.add_run(value)
        return

    p = cell.paragraphs[0]

    if not p.runs:
        p.add_run(value)
    else:
        p.runs[0].text = value
        for r in p.runs[1:]:
            r.text = ""


def fill_table_rows(table, trips, platform, receipt):
    """
    通用明细表填充：
    - 用 PLATFORM_SPECS[platform]["column_map"] 决定每列内容。
    - 假设第2行是占位模板行。
    """
    spec = PLATFORM_SPECS[platform]
    col_map = spec["column_map"]

    template_row = table.rows[1]

    # 清掉占位行之后已有的行
    for row in list(table.rows)[2:]:
        table._tbl.remove(row._tr)

    for idx, trip in enumerate(trips, start=1):
        new_tr = deepcopy(template_row._tr)
        table._tbl.append(new_tr)
        new_row = table.rows[-1]
        cells = new_row.cells

        for col_idx, fn in col_map.items():
            # 给列映射函数更多信息：trip, idx, receipt
            text = fn(trip, idx, receipt)
            set_cell(cells[col_idx], text)

    # 删除模板占位行自身
    table._tbl.remove(template_row._tr)


def build_docx_for_receipt(receipt, idx):
    platform = receipt.get("platform")
    spec = PLATFORM_SPECS[platform]

    doc = Document(spec["template"])

    # 1) 表头占位符替换（这版不会动 logo）
    fill_header_placeholders(doc, receipt, platform)

    # 2) 明细区域
    details_table = doc.tables[spec["details_table_index"]]

    if "row_block" in spec:
        fill_table_row_block(details_table, receipt["trips"], spec, receipt)
    else:
        # 走你已有的“单行模板”逻辑
        fill_table_rows(details_table, receipt["trips"], platform, receipt)

    # 3) 保存
    base_id = receipt.get("id") or f"{platform}_{idx:06d}"
    out_path = os.path.join(OUT_DOCX_DIR, f"{base_id}.docx")
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    return out_path


# def main():
#     with open(DATA_PATH, "r", encoding="utf-8") as f:
#         receipts = json.load(f)

#     # 先生成几条看效果
#     for idx, receipt in enumerate(receipts, start=1):
#         path = build_docx_for_receipt(receipt, idx)
#         print("生成：", path)
#         if idx >= 3:
#             break  # 先看前三条是否对版，没问题再去掉这个限制

# def main():
#     with open(DATA_PATH, "r", encoding="utf-8") as f:
#         receipts = json.load(f)

#     # 1) 按平台分组
#     by_platform = {}
#     for r in receipts:
#         p = r.get("platform", "baidu")
#         by_platform.setdefault(p, []).append(r)

#     # 2) 每个平台各生成 3 份样例 doc
#     for platform, items in by_platform.items():
#         print(f"--- 平台: {platform}，共有 {len(items)} 条，生成 3 份样例 ---")
#         subset = items[:3] if len(items) >= 3 else items  # 不足 3 条就全用

#         for i, receipt in enumerate(subset, start=1):
#             # 确保有唯一 id（不覆盖原 id）
#             base_id = receipt.get("id") or f"{platform}_{i:06d}"
#             r_copy = dict(receipt)
#             r_copy["id"] = base_id  # 调试期保留干净的命名

#             path = build_docx_for_receipt(r_copy, i)
#             print("生成：", path)



def main(force_regen: bool = False):
    """
    force_regen = False：增量模式，只给尚未有 docx 的样本生成 docx
    force_regen = True：全量重建，所有 id 的 docx 都重新生成（覆盖旧文件）
    """
    # 1) 读所有 faker 数据
    with open(DATA_PATH, "r", encoding="utf-8") as f:
        receipts = json.load(f)

    os.makedirs(OUT_DOCX_DIR, exist_ok=True)

    # 2) 收集一下已有的 docx 文件，方便日志 & 检查
    existing_docx = set(
        os.path.splitext(name)[0]
        for name in os.listdir(OUT_DOCX_DIR)
        if name.lower().endswith(".docx")
    )

    print(f"[INFO] 当前已有 DOCX 数量: {len(existing_docx)}")

    # 3) 逐条遍历 receipts，按 id 生成 docx
    for idx, receipt in enumerate(receipts, start=1):
        # if idx > 5:     # 🔴 只先跑前 3 条，调试用
        #     break       # 🔴 跑通后把这两行删掉即可
        rid = receipt.get("id")
        if not rid:
            print(f"[WARN] 第 {idx} 条没有 id，跳过")
            continue

        docx_name = f"{rid}.docx"
        docx_path = os.path.join(OUT_DOCX_DIR, docx_name)

        # ---- 增量模式：如果已经有 docx 且不强制重建，则跳过 ----
        if (not force_regen) and os.path.exists(docx_path):
            print(f"[SKIP] 已存在，跳过: {docx_path}")
            continue

        # 复制一份 receipt，避免修改原始数据
        r_copy = dict(receipt)
        r_copy["id"] = rid  # 确保 id 是我们想要的

        # build_docx_for_receipt 内部用 r_copy["id"] 来命名文件
        # 建议你在 build_docx_for_receipt 里用 id 当文件名，而不是 idx
        out_path = build_docx_for_receipt(r_copy, idx)

        print(f"[OK] 生成: {out_path}")

    print("[DONE] DOCX 生成流程结束。")

if __name__ == "__main__":
    main()